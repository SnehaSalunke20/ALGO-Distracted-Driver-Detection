from typing import Any, Dict
from pathlib import Path
import json
import base64
import os
import yaml
from yaml.loader import SafeLoader

import streamlit as st
import toml
import pandas as pd
import docx


#@st.cache(allow_output_mutation=True, ttl=300)
def get_project_root() -> str:
    """Returns project root path.

    Returns
    -------
    str
        Project root path.
    """
    return str(Path(__file__).parent.parent)

def load_auth_config(auth_config: str) -> Dict[Any, Any]:
    """Loads auth configuration file.

    Parameters
    ----------
    auth_config : str
        Filename of auth configuration file.

    Returns
    -------
    dict
        auth configuration file.
    """   
    file_path = Path(get_project_root()) / f"config/{auth_config}"
    with file_path.open() as file:
        config = yaml.load(file, Loader=SafeLoader)
    return config

def load_config(
    config_streamlit_filename: str) -> Dict[Any, Any]:
    """Loads configuration files.

    Parameters
    ----------
    config_streamlit_filename : str
        Filename of lib configuration file.

    Returns
    -------
    dict
        Lib configuration file.
    """
    config_streamlit = toml.load(Path(get_project_root()) / f"config/{config_streamlit_filename}")
    return dict(config_streamlit)

# get key metrics data
def get_key_metrics_data(news_for):
    file_path_stats_valuation = f"data/stats_valuation_data/{news_for}.json"
    with open(file_path_stats_valuation, "r") as json_file:
        key_metrics_data = json.load(json_file)
    return key_metrics_data


# if __name__ == '__main__':
#     print(load_config("config_streamlit.toml"))

# load lottie files
@st.cache_resource 
def load_lottiefile(filepath: str):
    with open(filepath, "r") as f:
        return json.load(f)
    
@st.cache_resource    
def load_base64_image(image_path):
    with open(image_path, "rb") as img_file:
            base64_image = base64.b64encode(img_file.read()).decode()
            
    return base64_image


def displayPDF(file_or_buffer, width=700, height=750):
    # Check if the input is a file path or a BytesIO object
    if isinstance(file_or_buffer, str):  # file path
        with open(file_or_buffer, "rb") as f:
            file_content = f.read()
    else:  # BytesIO object
        file_content = file_or_buffer.read()

    base64_pdf = base64.b64encode(file_content).decode('utf-8')

    # Embedding PDF in HTML
    pdf_display = F'<embed src="data:application/pdf;base64,{base64_pdf}" width=100% height="{height}" type="application/pdf">'
    # Displaying File
    st.markdown(pdf_display, unsafe_allow_html=True)




def displayFile(file_path_or_stream):
    # Determine if input is a file path or a stream
    if isinstance(file_path_or_stream, str):  
        extension = os.path.splitext(file_path_or_stream)[-1].lower()
    else:
        extension = "." + file_path_or_stream.name.split('.')[-1].lower()

    # Display based on file extension
    if extension == ".pdf":
        displayPDF(file_path_or_stream)
    elif extension == ".txt":
        if isinstance(file_path_or_stream, str):
            with open(file_path_or_stream, 'r') as file:
                content = file.read()
        else:
            content = file_path_or_stream.read().decode('utf-8')
        st.text_area("Text Content", content, height=400)
    elif extension == ".docx":
        if isinstance(file_path_or_stream, str):
            doc = docx.Document(file_path_or_stream)
        else:
            file_path_or_stream.seek(0)
            doc = docx.Document(file_path_or_stream)
        for paragraph in doc.paragraphs:
            st.write(paragraph.text)
    elif extension == ".xlsx":
        if isinstance(file_path_or_stream, str):
            df = pd.read_excel(file_path_or_stream)
        else:
            file_path_or_stream.seek(0)
            df = pd.read_excel(file_path_or_stream)
        st.write(df)
    else:
        st.warning(f"File type {extension} not supported for display!")









def custom_footer(_gettext):

    # Load base64 images
    company_logo_base64 = load_base64_image("ui_assets/images/algo-logo.png")
    ig_image = load_base64_image("ui_assets/images/instagram-logo-transparent.png")
    linked_in_image = load_base64_image("ui_assets/images/linkedin-logo-transparent.png")
    fb_image = load_base64_image("ui_assets/images/facebook-logo-png.png")


    st.markdown("---")
    footer_html_styles = """
    <style>
    body {
         /* Ensure the body is relatively positioned */
        margin-bottom: 0px; /* Adjust this value to match the desired footer height */
    }
    .footer-container {
        position: absolute;
        bottom: 5000;
        width: 100%;
        margin-bottom: 10px; /* Adjust this value to move the footer up */
        background-color: #f9f9f9;
        color: #6c757d;
        font-family: 'Arial', sans-serif;
        display: flex;
        justify-content: space-between;
        align-items: center;
        font-size: 11px;
        border-bottom: 1px solid #d3d3d3;
    }

    .footer {
        width: 100%;
        background-color: #789be6;
        color: white;
        display: grid;
        grid-template-columns: repeat(1, 1fr);
        gap: 2rem;
        padding: 3rem 2rem;
      }

    .contact-section {
        display: flex;
        flex-direction: column;
        text-align: center;
        align-items: center;
      }

      .contact-text {
        font-size: 1.5rem;
        line-height: 1.5rem;
        margin-bottom: 1.5rem;
        max-width: 700px;
      }

      

      .contact-address {
        font-size: 0.875rem;
        line-height: 1.25rem;
      }

      .links-section {
        margin-top: 2rem;
        display: flex;
        flex-direction: column;
        align-items: center;
        gap: 2rem;
      }

      .social-button {
        background-color: white;
        border: none;
        border-radius: 0.3rem;
        padding: 0.5rem 1.25rem;
        color: black;
        display: inline-flex;
        align-items: center;
        font-size: 1rem;
        line-height: 1.5rem;
      }

      .social-button img {
        height: 1rem;
        width: 1rem;
        margin-left: 0.5rem;
      }

      .text-links {
        font-size: 0.875rem;
        line-height: 1.25rem;
        text-align: center;
      }

      .text-links a {
        display: block;
        margin-top: 0.5rem;
        color: inherit;
        text-decoration: inherit;
      }

      .social-links {
        display: flex;
        align-items: center;
        gap: 1rem;
        justify-content: center;
      }
      .social-links a {
        padding: 0.15rem;
        border-radius: 1rem;
        background-color: white;
      }
      .social-links a img {
        height: 2rem;
        width: 2rem;
        margin: 0.5rem;
      }

      .bottom-section {
        display: flex;
        flex-direction: column;
        align-items: center;
        border-top: 1px solid black;
        grid-column: 1 / -1;
        padding: 0.5rem;
      }

      .bottom-section a {
        text-decoration: none !important;
      }

      .logo-link {
        display: flex;
        align-items: center;
        gap: 1rem;
        text-decoration: none;
        color: inherit;
      }

      .logo {
        height: 4rem;
        width: 4rem;
      }

      .company-name {
        color: white;
        text-decoration: none;
        font-size: 1.125rem;
        line-height: 1.75rem;
      }

      .copyright-info {
        display: flex;
        align-items: center;
        gap: 2rem;
        margin-left: auto;
        font-size: 0.75rem;
        line-height: 1rem;
      }

      .copyright-info a {
        color: inherit;
        text-decoration: none;
      }

      @media (min-width: 768px) {
        .footer {
          grid-template-columns: repeat(2, 1fr);
          padding: 2rem 5rem;
        }
        .contact-section {
          text-align: left;
          align-items: flex-start;
        }
        .contact-text {
          font-size: 1.75rem;
          line-height: 2.5rem;
        }
        .contact-address {
          font-size: 1rem;
          line-height: 1.5rem;
        }
        .links-section {
          flex-direction: row;
          align-items: flex-end;
        }
        .text-links {
          text-align: left;
          font-size: 1rem;
          line-height: 1.5rem;
        }
        .text-links {
          text-align: left;
          font-size: 1rem;
          line-height: 1.5rem;
        }
        .bottom-section {
          flex-direction: row;
          justify-content: space-between;
          align-items: center;
        }
        .company-name {
          font-size: 1.25rem;
        }
        .copyright-info {
          font-size: 0.875rem;
          line-height: 1.25rem;
        }
      }
    </style>
    """
    line1 = _gettext("Let's Connect!")
    line2 = _gettext("We'd love to hear from you.")
    mail = _gettext("Mail Us")
    privacy = _gettext("Privacy Policy")
    company = _gettext("AlgoAnalytics")
    add = _gettext("4th Floor, Algoanalytics Pvt. Ltd., Alacrity India Innovation Centre, Ideas to Impacts Building. Pallod Farm Lane 3, Near Vijay Sales, Baner Road, Pune - 411045")

    footer_html_body = f"""
    <div style = "margin-bottom: 0; margin-top:0;"class="footer-container">
    <div class="footer">
        <div class="contact-section">
            <div class="contact-text">
                    {line1}
                    <br />
                    {line2}
            </div>
            <div class="social-links">
            <a
              href="https://www.facebook.com/people/Algoanalytics/100069650841271/"
              target="_blank"
              aria-label="Facebook"
            >
             <img src='data:image/png;base64,{fb_image}'/>
            </a>
            <a
              href="https://www.instagram.com/algoanalyticsin/"
              target="_blank"
              aria-label="Instagram"
            >
             <img src='data:image/png;base64,{ig_image}'/>
            </a>
            <a
              href="https://www.linkedin.com/company/algoanalytics/"
              target="_blank"
              aria-label="LinkedIn"
            >
              <img src='data:image/png;base64,{linked_in_image}'/>
            </a>
            <a
              href="https://goo.gl/maps/gFbb2TYKHSGfJSi17"
              target="_blank"
              aria-label="Maps"
            >
              <img src='data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIwLjdlbSIgaGVpZ2h0PSIxZW0iIHZpZXdCb3g9IjAgMCAyNTYgMzY3Ij48cGF0aCBmaWxsPSIjMzRhODUzIiBkPSJNNzAuNTg1IDI3MS44NjVhMzcwLjcxMiAzNzAuNzEyIDAgMCAxIDI4LjkxMSA0Mi42NDJjNy4zNzQgMTMuOTgyIDEwLjQ0OCAyMy40NjMgMTUuODM3IDQwLjMxYzMuMzA1IDkuMzA4IDYuMjkyIDEyLjA4NiAxMi43MTQgMTIuMDg2YzYuOTk4IDAgMTAuMTczLTQuNzI2IDEyLjYyNi0xMi4wMzVjNS4wOTQtMTUuOTEgOS4wOTEtMjguMDUyIDE1LjM5Ny0zOS41MjVjMTIuMzc0LTIyLjE1IDI3Ljc1LTQxLjgzMyA0Mi44NTgtNjAuNzVjNC4wOS01LjM1NCAzMC41MzQtMzYuNTQ1IDQyLjQzOS02MS4xNTZjMCAwIDE0LjYzMi0yNy4wMzUgMTQuNjMyLTY0Ljc5MmMwLTM1LjMxOC0xNC40My01OS44MTMtMTQuNDMtNTkuODEzbC00MS41NDUgMTEuMTI2bC0yNS4yMyA2Ni40NTFsLTYuMjQyIDkuMTYzbC0xLjI0OCAxLjY2bC0xLjY2IDIuMDc4bC0yLjkxNCAzLjMxOWwtNC4xNjQgNC4xNjNsLTIyLjQ2NyAxOC4zMDRsLTU2LjE3IDMyLjQzMnoiLz48cGF0aCBmaWxsPSIjZmJiYzA0IiBkPSJNMTIuNjEyIDE4OC44OTJjMTMuNzA5IDMxLjMxMyA0MC4xNDUgNTguODM5IDU4LjAzMSA4Mi45OTVsOTUuMDAxLTExMi41MzRzLTEzLjM4NCAxNy41MDQtMzcuNjYyIDE3LjUwNGMtMjcuMDQzIDAtNDguODktMjEuNTk1LTQ4Ljg5LTQ4LjgyNWMwLTE4LjY3MyAxMS4yMzQtMzEuNTAxIDExLjIzNC0zMS41MDFsLTY0LjQ4OSAxNy4yOHoiLz48cGF0aCBmaWxsPSIjNDI4NWY0IiBkPSJNMTY2LjcwNSA1Ljc4N2MzMS41NTIgMTAuMTczIDU4LjU1OCAzMS41MyA3NC44OTMgNjMuMDIzbC03NS45MjUgOTAuNDc4czExLjIzNC0xMy4wNiAxMS4yMzQtMzEuNjE3YzAtMjcuODY0LTIzLjQ2My00OC42OC00OC44MS00OC42OGMtMjMuOTY5IDAtMzcuNzM1IDE3LjQ3NS0zNy43MzUgMTcuNDc1di01N3oiLz48cGF0aCBmaWxsPSIjMWE3M2U4IiBkPSJNMzAuMDE1IDQ1Ljc2NUM0OC44NiAyMy4yMTggODIuMDIgMCAxMjcuNzM2IDBjMjIuMTggMCAzOC44OSA1LjgyMyAzOC44OSA1LjgyM0w5MC4yOSA5Ni41MTZIMzYuMjA1eiIvPjxwYXRoIGZpbGw9IiNlYTQzMzUiIGQ9Ik0xMi42MTIgMTg4Ljg5MlMwIDE2NC4xOTQgMCAxMjguNDE0YzAtMzMuODE3IDEzLjE0Ni02My4zNzcgMzAuMDE1LTgyLjY0OWw2MC4zMTggNTAuNzU5eiIvPjwvc3ZnPg=='/>
            </a>
          </div>
        </div>
        <div class="">
           <div class="contact-address">
                {add}
            </div>
            <div class="links-section">
              <a href="mailto:info@algoanalytics.com">
                  <button class="social-button">
                  {mail}
                  <img src='data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxLjc4ZW0iIGhlaWdodD0iMWVtIiB2aWV3Qm94PSIwIDAgMTYgOSI+PHBhdGggZmlsbD0iY3VycmVudENvbG9yIiBkPSJNMTIuNSA1aC05Yy0uMjggMC0uNS0uMjItLjUtLjVzLjIyLS41LjUtLjVoOWMuMjggMCAuNS4yMi41LjVzLS4yMi41LS41LjUiLz48cGF0aCBmaWxsPSJjdXJyZW50Q29sb3IiIGQ9Ik0xMCA4LjVhLjQ3LjQ3IDAgMCAxLS4zNS0uMTVjLS4yLS4yLS4yLS41MSAwLS43MWwzLjE1LTMuMTVsLTMuMTUtMy4xNWMtLjItLjItLjItLjUxIDAtLjcxcy41MS0uMi43MSAwbDMuNSAzLjVjLjIuMi4yLjUxIDAgLjcxbC0zLjUgMy41Yy0uMS4xLS4yMy4xNS0uMzUuMTVaIi8+PC9zdmc+' alt='right-arrow'/>
                  </button>
              </a>
              <a href="https://onestop.ai/Privacy-Policy">
                  <button class="social-button">
                  {privacy}
                  <img src='data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxZW0iIGhlaWdodD0iMWVtIiB2aWV3Qm94PSIwIDAgMjggMjgiPjxnIGZpbGw9ImN1cnJlbnRDb2xvciI+PHBhdGggZD0iTTE0IDJ2OGEyIDIgMCAwIDAgMiAyaDcuOTk5bC4wMDEuMDc4VjIzLjZhMi40IDIuNCAwIDAgMS0yLjQgMi40SDYuNEEyLjQgMi40IDAgMCAxIDQgMjMuNlY0LjRBMi40IDIuNCAwIDAgMSA2LjQgMnoiLz48cGF0aCBkPSJNMTUuNSAyLjQ3NVYxMGEuNS41IDAgMCAwIC41LjVoNy41MDJhMyAzIDAgMCAwLS4zMDctLjM2NmwtNy40MzEtNy40MzFhMi40IDIuNCAwIDAgMC0uMjY0LS4yMjgiLz48L2c+PC9zdmc+' alt='right-arrow'/>
                  </button>
              </a>
            </div>
        </div>
        <div class="bottom-section">
          <a href="https://algoanalytics.com/" target="_blank">
            <div class="logo-link">
              <img
                class="logo"
                src='data:image/png;base64,{company_logo_base64}'
                alt="Company Logo"
              />
              <div class="company-name">{company}</div>
            </div>
          </a>
          <div class="copyright-info">
            <span class="font-thin">Onestop.ai© 2023 </span>
            <a href="mailto:info@algoanalytics.com">info@algoanalytics.com</a>
          </div>
        </div>
    </div>
    </div>
    """
    
    st.markdown(footer_html_styles + footer_html_body, unsafe_allow_html=True)

   

def apply_custom_style():
    st.markdown("""
        <style>
            .block-container {
                padding-top: 0rem !important;
                padding-left: 1rem;
                padding-right: 1rem;
            }
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
        </style>
    """, unsafe_allow_html=True)
